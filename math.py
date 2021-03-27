#coding=utf-8
 
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import qn
import random

def genMath(max,op=None):
    if not op:
        op = random.choice(['+','-'])
    if op=='+':
        a = random.randint(0,max/2)
        b = random.randint(0,max-a)
        return ("%s + %s = " % (a,b))
    if op=='-':
        a = random.randint(5,max)
        b = random.randint(0,a)
        return ("%s - %s = " % (a,b))

def genMath10(op=None):
    if not op:
        op = random.choice(['+','-'])
    if op=='+':
        a = random.randint(0,10)
        b = 10-a
        return ("%s + %s = " % (a,b))
    if op=='-':
        a = 10
        b = random.randint(0,10)
        return ("%s - %s = " % (a,b))

#打开文档
doc = Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.styles['Normal'].font.size=Pt(20)
doc.styles['Normal'].font.bold = True
'''
#加入不同等级的标题
document.add_heading(u'MS WORD写入测试',0)
document.add_heading(u'一级标题',1)
document.add_heading(u'二级标题',2)
#添加文本
paragraph = document.add_paragraph(u'我们在做文本测试！')
#设置字号
run = paragraph.add_run(u'设置字号、')
run.font.size = Pt(24)
 
#设置字体
run = paragraph.add_run('Set Font,')
run.font.name = 'Consolas'
 
#设置中文字体
run = paragraph.add_run(u'设置中文字体、')
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
 
#设置斜体
run = paragraph.add_run(u'斜体、')
run.italic = True
 
#设置粗体
run = paragraph.add_run(u'粗体').bold = True
 
#增加引用
document.add_paragraph('Intense quote', style='Intense Quote')
 
#增加无序列表
document.add_paragraph(
    u'无序列表元素1', style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2', style='List Bullet'
)
#增加有序列表
document.add_paragraph(
    u'有序列表元素1', style='List Number'
)
document.add_paragraph(
    u'有序列表元素2', style='List Number'
)
#增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
#document.add_picture('image.bmp', width=Inches(1.25))
'''
#增加表格
table = doc.add_table(rows=15, cols=2)
'''
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
'''
#再增加3行表格元素
for row in table.rows:
    for cell in row.cells:
        cell.text = genMath(30)

#增加分页
#doc.add_page_break()
 
#保存文件
doc.save(u'测试.docx')